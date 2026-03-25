from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

import pytest
from docx import Document
from fastapi import UploadFile
from starlette.datastructures import Headers

from memory_module.chunking.data_models import Chunk, ChunkMetadata
from memory_module.parser.data_models import DocumentParserResult, FileMetadata, ParsedContent, ParsedSection


def make_upload_file(
    filename: str,
    content: bytes,
    content_type: str,
) -> UploadFile:
    return UploadFile(
        file=BytesIO(content),
        size=len(content),
        filename=filename,
        headers=Headers({"content-type": content_type}),
    )


def build_docx_bytes(paragraphs: list[str | tuple[str, str]]) -> bytes:
    document = Document()
    for paragraph in paragraphs:
        if isinstance(paragraph, tuple):
            text, style = paragraph
            document.add_paragraph(text, style=style)
        else:
            document.add_paragraph(paragraph)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


@pytest.fixture
def docx_bytes() -> bytes:
    return build_docx_bytes(["First paragraph", "Second paragraph"])


@pytest.fixture
def upload_docx(docx_bytes: bytes) -> UploadFile:
    return make_upload_file(
        filename="sample.docx",
        content=docx_bytes,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


@pytest.fixture
def parsed_document() -> DocumentParserResult:
    return DocumentParserResult(
        content=ParsedContent(
            mode="text",
            text="alpha beta gamma delta",
            sections=[
                ParsedSection(title="Intro", text="alpha beta"),
                ParsedSection(title="Body", text="gamma delta"),
            ],
        ),
        file_metadata=FileMetadata(
            document_id="doc_test",
            document_title="Sample Doc",
        ),
    )


@pytest.fixture
def sample_chunk() -> Chunk:
    return Chunk(
        chunk_id="chunk-1",
        text="alpha beta",
        embedding=[0.1, 0.2, 0.3],
        metadata=ChunkMetadata(
            document_id="doc_test",
            document_title="Sample Doc",
            tags=["tag-a"],
            chunk_version="doc_test_chunk_1",
        ),
        token_count=2,
    )


class StubParser:
    def __init__(self, parsed_document: DocumentParserResult, accepts: bool = True):
        self._parsed_document = parsed_document
        self._accepts = accepts
        self.last_error: str | None = None
        self.calls: list[str] = []

    def accepts(self, file_stream: UploadFile) -> bool:
        self.calls.append("accepts")
        if not self._accepts:
            self.last_error = "Rejected for test"
        return self._accepts

    def convert(self, file_stream: UploadFile) -> DocumentParserResult:
        self.calls.append("convert")
        return self._parsed_document


class StubChunker:
    def __init__(self, chunks: list[Chunk]):
        self._chunks = chunks
        self.calls: list[tuple[str, Any]] = []

    def chunk(self, parsed_document: DocumentParserResult, metadata: dict[str, Any]) -> list[Chunk]:
        self.calls.append(("chunk", parsed_document, metadata))
        return self._chunks


class StubEmbedder:
    def __init__(self, embedding: list[float] | list[list[float]]):
        self.embedding = embedding
        self.calls: list[str] = []

    def embed(self, text: str):
        self.calls.append(text)
        return self.embedding


class StubVectorDB:
    def __init__(self):
        self.added_chunks: list[Chunk] | None = None

    def add_chunks(self, chunks: list[Chunk]):
        self.added_chunks = chunks
